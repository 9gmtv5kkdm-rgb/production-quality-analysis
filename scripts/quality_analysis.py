#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
quality_analysis.py — 生产质量归因分析核心引擎（能力层落地脚本）

能力：读表 → 缺陷族归类 → 帕累托 + 工序×族交叉表 → 参数化鱼骨图/帕累托 SVG
      → 可选 DOCX 报告骨架。纯标准库生成 SVG，无需第三方绘图库。

依赖：
  - 读 xlsx：pandas + openpyxl（pip install pandas openpyxl）
  - 写 DOCX：python-docx（可选，缺则跳过，仅输出 HTML/SVG/JSON）
运行：
  python quality_analysis.py data.xlsx --desc 异常描述 --qty 数量 --proc 工序 \
         --html out.html --docx out.docx
不传参时打印本模块自检（各函数自测），证明能力可用。
"""
import argparse
import json
import sys
from collections import OrderedDict

# ---------------------------------------------------------------------------
# 1. 缺陷族分类（与 references/defect-taxonomy.md 一致，可行业替换）
# ---------------------------------------------------------------------------
DEFECT_FAMILIES = OrderedDict([
    ("H", r"裂纹|开裂|碎脆|母材|铁素体|固溶|性能"),
    ("G", r"错边|台阶|焊缝|未熔合|焊"),
    ("F", r"坡口|修磨|削薄|打磨|外削薄|内削薄|减薄"),
    ("E", r"划伤|旋痕|凹坑|麻点|压痕|粘模|氧化皮|拉丝|伤|痕"),
    ("D", r"腰线|花边|鼓包|凸起|压坑|变形|褶皱|棱"),
    ("C", r"外倾|内倾|外翻|直边"),
    ("B", r"圆度|椭圆|同心度"),
    ("I", r"高度|深度"),
    ("A", r"直径|周长|内径|外径|口径|尺寸|偏大|偏小|超差|一致|余量"),
])
# 族码 → 中文名
FAMILY_NAME = {
    "A": "直径/尺寸", "B": "圆度/椭圆/同心度", "C": "直边外倾/内倾",
    "D": "腰线/花边/鼓包/变形", "E": "划伤/旋痕/凹坑/表面", "F": "坡口/修磨/减薄",
    "G": "错边/台阶/焊缝", "H": "裂纹/开裂/材质", "I": "高度/深度", "J": "其它",
}
# 工序归一化
PROC_MAP = [
    (r"下料|锯|切", "下料"), (r"CNC|车|铣|钻|精加工", "机加工"),
    (r"焊|组对|拼焊", "焊接"), (r"装|组装|总装", "装配"),
    (r"热|正火|淬火|固溶", "热处理"), (r"喷|涂|电泳", "喷涂"),
    (r"冲|压制|成型", "冲压"), (r"注|注塑|模压", "注塑"),
    (r"检|终检|三检", "检验"), (r"返|返修|补焊", "返修"),
]
import re

def classify_family(text):
    if not isinstance(text, str) or not text.strip():
        return "J"
    for code, pat in DEFECT_FAMILIES.items():
        if re.search(pat, text, re.IGNORECASE):
            return code
    return "J"

def normalize_process(text):
    if not isinstance(text, str) or not text.strip():
        return "未记录"
    for pat, name in PROC_MAP:
        if re.search(pat, text, re.IGNORECASE):
            return name
    return text.strip()

# ---------------------------------------------------------------------------
# 2. 读表
# ---------------------------------------------------------------------------
def read_table(path, header_row=1, desc_col="异常描述", qty_col="数量", proc_col="工序"):
    """返回 DataFrame，含归一化列 defect_family / process_norm / qty。"""
    try:
        import pandas as pd
    except ImportError:
        raise SystemExit("需要 pandas+openpyxl：pip install pandas openpyxl")
    df = pd.read_excel(path, header=header_row)
    df.columns = [str(c).strip() for c in df.columns]
    # 容错：列名包含关键字即可
    def find_col(hint):
        for c in df.columns:
            if hint in c:
                return c
        return hint
    dcol = find_col(desc_col)
    qcol = find_col(qty_col)
    pcol = find_col(proc_col) if proc_col in df.columns or any(proc_col in c for c in df.columns) else None
    df = df[df[dcol].notna()]
    df["defect_family"] = df[dcol].apply(classify_family)
    df["process_norm"] = (df[pcol].apply(normalize_process) if pcol else "未记录")
    df["qty"] = pd.to_numeric(df[qcol], errors="coerce").fillna(1)
    return df

# ---------------------------------------------------------------------------
# 3. 帕累托 + 交叉表
# ---------------------------------------------------------------------------
def pareto(df, by="defect_family"):
    g = df.groupby(by)["qty"].agg(["sum", "count"]).reset_index()
    g = g.sort_values("sum", ascending=False)
    total = g["sum"].sum()
    g["pct"] = g["sum"] / total
    g["cum"] = g["pct"].cumsum()
    # 关键少数判定
    def grade(c):
        if c <= 0.80: return "关键"
        if c <= 0.95: return "次要"
        return "一般"
    g["grade"] = g["cum"].apply(grade)
    g[by] = g[by].map(lambda c: f"{c} {FAMILY_NAME.get(c, c)}")
    return g, total

def cross_tab(df):
    return df.pivot_table(index="process_norm", columns="defect_family",
                          values="qty", aggfunc="sum", fill_value=0)

# ---------------------------------------------------------------------------
# 4. SVG 生成（标准库，无依赖）
# ---------------------------------------------------------------------------
def _esc(s):
    return (str(s).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"))

def fishbone_svg(effect, cats, width=1040, height=620):
    """cats: dict 维度名 -> list[原因str]。3 上 3 下斜肋。"""
    order = list(cats.keys())
    above = order[: (len(order) + 1) // 2]
    below = order[len(above):]
    mid_y = height // 2
    spine_x0, spine_x1 = 90, width - 210
    parts = [f'<svg xmlns="http://www.w3.org/2000/svg" width="{width}" height="{height}" '
             f'viewBox="0 0 {width} {height}" font-family="Microsoft YaHei, sans-serif" font-size="13">']
    # 脊柱 + 箭头
    parts.append(f'<line x1="{spine_x0}" y1="{mid_y}" x2="{spine_x1+30}" y2="{mid_y}" '
                 f'stroke="#2C5AA0" stroke-width="3"/>')
    parts.append(f'<polygon points="{spine_x1+30},{mid_y} {spine_x1+12},{mid_y-8} {spine_x1+12},{mid_y+8}" '
                 f'fill="#2C5AA0"/>')
    # 效应框
    parts.append(f'<rect x="{spine_x1+34}" y="{mid_y-26}" width="170" height="52" rx="8" '
                 f'fill="#2C5AA0"/>')
    parts.append(f'<text x="{spine_x1+119}" y="{mid_y+5}" fill="#fff" text-anchor="middle" '
                 f'font-weight="bold">{_esc(effect)}</text>')

    def bone(cat, idx, side):  # side: -1 上, +1 下
        n = len(above) if side < 0 else len(below)
        i = above.index(cat) if side < 0 else below.index(cat)
        ax = spine_x1 - 40 - i * ((spine_x1 - spine_x0 - 80) / max(n, 1))
        ay = mid_y
        lx = ax - 70
        ly = mid_y + side * (90 + i * 70)
        parts.append(f'<line x1="{ax}" y1="{ay}" x2="{lx}" y2="{ly}" stroke="#2C5AA0" stroke-width="2"/>')
        parts.append(f'<text x="{lx-8}" y="{ly+ (4 if side<0 else 14)}" text-anchor="end" '
                     f'font-weight="bold" fill="#1A3A6B">{_esc(cat)}</text>')
        # 子原因沿肋排列
        for j, cause in enumerate(cats[cat]):
            ty = ly + side * (j + 1) * 18
            parts.append(f'<text x="{lx+6}" y="{ty}" fill="#333">{_esc(cause)}</text>')

    for cat in above:
        bone(cat, 0, -1)
    for cat in below:
        bone(cat, 0, 1)
    parts.append('</svg>')
    return "\n".join(parts)

def pareto_svg(rows, width=900, height=460):
    """rows: list of (label, qty, pct, cum)。柱+累计折线+80%参考线。"""
    if not rows:
        return "<svg/>"
    max_q = max(r[1] for r in rows)
    n = len(rows)
    bw = (width - 120) / n
    base_y = height - 60
    parts = [f'<svg xmlns="http://www.w3.org/2000/svg" width="{width}" height="{height}" '
             f'viewBox="0 0 {width} {height}" font-family="Microsoft YaHei, sans-serif" font-size="12">']
    # 80% 参考线
    y80 = base_y - 0.8 * (base_y - 50)
    parts.append(f'<line x1="100" y1="{y80:.0f}" x2="{width-20}" y2="{y80:.0f}" '
                 f'stroke="#E08214" stroke-dasharray="6 4"/>')
    parts.append(f'<text x="{width-18}" y="{y80-6:.0f}" fill="#E08214" text-anchor="end">80%</text>')
    for i, (label, qty, pct, cum) in enumerate(rows):
        x = 100 + i * bw
        h = (qty / max_q) * (base_y - 50)
        parts.append(f'<rect x="{x+4:.0f}" y="{base_y-h:.0f}" width="{bw-8:.0f}" height="{h:.0f}" '
                     f'fill="#2C5AA0"/>')
        parts.append(f'<text x="{x+bw/2:.0f}" y="{base_y+14}" fill="#333" text-anchor="middle" '
                     f'transform="rotate(30 {x+bw/2:.0f} {base_y+14})">{_esc(label)}</text>')
        parts.append(f'<text x="{x+bw/2:.0f}" y="{base_y-h-6:.0f}" fill="#2C5AA0" '
                     f'text-anchor="middle">{qty}</text>')
        # 累计点
        cy = base_y - cum * (base_y - 50)
        parts.append(f'<circle cx="{x+bw/2:.0f}" cy="{cy:.0f}" r="3" fill="#E08214"/>')
        if i > 0:
            px = 100 + (i - 1) * bw + bw / 2
            py = base_y - rows[i - 1][3] * (base_y - 50)
            parts.append(f'<line x1="{px:.0f}" y1="{py:.0f}" x2="{x+bw/2:.0f}" y2="{cy:.0f}" '
                         f'stroke="#E08214" stroke-width="2"/>')
    parts.append('</svg>')
    return "\n".join(parts)

# ---------------------------------------------------------------------------
# 5. 可选 DOCX（python-docx）
# ---------------------------------------------------------------------------
def generate_docx(path, pareto_df, cross=None):
    try:
        from docx import Document
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        print("[warn] 未安装 python-docx，跳过 DOCX 生成（pip install python-docx）")
        return False
    doc = Document()
    doc.add_heading("生产质量归因分析报告", 0)
    t = doc.add_table(rows=1, cols=len(pareto_df.columns))
    t.style = "Light Grid Accent 1"
    for j, c in enumerate(pareto_df.columns):
        t.rows[0].cells[j].text = str(c)
    for _, row in pareto_df.iterrows():
        cells = t.add_row().cells
        for j, c in enumerate(pareto_df.columns):
            cells[j].text = str(row[c])
    doc.save(path)
    return True

# ---------------------------------------------------------------------------
# 6. CLI
# ---------------------------------------------------------------------------
def main():
    ap = argparse.ArgumentParser(description="生产质量归因分析引擎")
    ap.add_argument("path", help="xlsx 台账路径")
    ap.add_argument("--desc", default="异常描述")
    ap.add_argument("--qty", default="数量")
    ap.add_argument("--proc", default="工序")
    ap.add_argument("--header", type=int, default=1)
    ap.add_argument("--html", help="输出 HTML 报告")
    ap.add_argument("--docx", help="输出 DOCX 报告")
    ap.add_argument("--json", help="输出帕累托 JSON")
    args = ap.parse_args()

    df = read_table(args.path, args.header, args.desc, args.qty, args.proc)
    g, total = pareto(df)
    ct = cross_tab(df)
    print(f"总异常件数={int(total)}  记录数={len(df)}")
    print(g.to_string(index=False))
    print("\n工序×缺陷族交叉表：")
    print(ct.to_string())

    if args.json:
        g.to_json(args.json, orient="records", force_ascii=False, indent=2)
        print(f"\n[ok] 帕累托 JSON -> {args.json}")
    if args.html:
        svg_p = pareto_svg([(r.split()[0], r_q, r_p, r_c) for r, r_q, r_p, r_c in
                            zip(g[g.columns[0]], g["sum"], g["pct"], g["cum"])])
        with open(args.html, "w", encoding="utf-8") as f:
            f.write(f"<html><body><h1>质量归因</h1>{svg_p}<pre>{g.to_string(index=False)}</pre></body></html>")
        print(f"[ok] HTML -> {args.html}")
    if args.docx:
        generate_docx(args.docx, g)

if __name__ == "__main__":
    if len(sys.argv) == 1:
        # 自检：证明各能力函数可用
        print("== 自检 classify_family ==")
        for s in ["直径偏大超差", "表面划伤", "焊缝错边", "裂纹开裂", "未知问题"]:
            print(f"  {s!r} -> {classify_family(s)}")
        print("== 自检 fishbone_svg（前80字符）==")
        svg = fishbone_svg("尺寸超差", {"人": ["未按SOP"], "机": ["夹具磨损"],
                            "料": ["母材波动"], "法": ["首检缺失"], "环": ["温差"], "测": ["量具失准"]})
        print("  len=", len(svg), svg[:80])
        print("== 自检 pareto_svg ==")
        print("  len=", len(pareto_svg([("A", 320, .34, .34), ("E", 210, .22, .56)])))
        print("模块能力可用。")
    else:
        main()
